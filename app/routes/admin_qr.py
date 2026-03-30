# app/routes/admin_qr.py
"""
Admin routes for QR code generation.

Generates QR codes for each guest's personalized RSVP link.

Filtering: only guests who are pending (no RSVP yet) or confirmed
           (is_attending=True, is_cancelled=False) are included.

Async bulk download: ZIP of DOCX files is generated in a background thread
                     to avoid blocking Flask workers. Poll /download-all-docx/status/<job_id>
                     and download from /download-all-docx/download/<job_id> when done.
"""

import threading
import uuid
import zipfile
import logging
from datetime import datetime, timedelta
from functools import wraps
from io import BytesIO

import qrcode
import qrcode.image.svg
from flask import (
    Blueprint, render_template, request, send_file,
    redirect, url_for, current_app, jsonify
)

from app.constants import Security

logger = logging.getLogger(__name__)

bp = Blueprint('admin_qr', __name__, url_prefix='/admin/qr')


# ---------------------------------------------------------------------------
# Async job store
# In-memory dict is fine for a single-server wedding site.
# Each job: {status, progress, total, data (BytesIO), error, created_at}
# ---------------------------------------------------------------------------
_docx_jobs: dict = {}
_jobs_lock = threading.Lock()

_JOB_TTL_MINUTES = 10  # Auto-expire unretrieved jobs


def _cleanup_old_jobs() -> None:
    """Remove jobs older than _JOB_TTL_MINUTES. Call before creating new jobs."""
    cutoff = datetime.utcnow() - timedelta(minutes=_JOB_TTL_MINUTES)
    with _jobs_lock:
        expired = [jid for jid, job in _docx_jobs.items()
                   if job["created_at"] < cutoff]
        for jid in expired:
            del _docx_jobs[jid]
    if expired:
        logger.info(f"Cleaned up {len(expired)} expired DOCX jobs.")


# ---------------------------------------------------------------------------
# Auth decorator
# ---------------------------------------------------------------------------

def admin_required(f):
    """Decorator to require admin authentication."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not request.cookies.get(Security.ADMIN_COOKIE_NAME):
            return redirect(url_for('admin.login'))
        return f(*args, **kwargs)
    return decorated_function


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _filter_guests(all_guests: list) -> list:
    """
    Return only guests who should receive a QR code:
      - Pending:   no RSVP submitted yet (guest.rsvp is None)
      - Confirmed: RSVP submitted, is_attending=True, is_cancelled=False
    """
    return [
        g for g in all_guests
        if g.rsvp is None or (g.rsvp.is_attending and not g.rsvp.is_cancelled)
    ]


def _safe_filename(name: str, surname) -> str:
    """Build a filesystem-safe filename stem from name + surname."""
    full = f"{name} {surname}" if surname else name
    safe = "".join(c for c in full if c.isalnum() or c in (" ", "-", "_")).strip()
    return safe.replace(" ", "_")


def generate_qr_code(url: str, format: str = "png", size: int = 10) -> BytesIO:
    """
    Generate a QR code image for a URL.

    Args:
        url:    The URL to encode.
        format: 'png' or 'svg'.
        size:   Box size in pixels (PNG only).

    Returns:
        BytesIO buffer positioned at offset 0.
    """
    if format == "svg":
        factory = qrcode.image.svg.SvgPathImage
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=size,
            border=2,
            image_factory=factory,
        )
        qr.add_data(url)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        buf = BytesIO()
        img.save(buf)
        buf.seek(0)
        return buf
    else:
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=size,
            border=2,
        )
        qr.add_data(url)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        buf = BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf


def _generate_guest_docx(name: str, surname, rsvp_url: str) -> BytesIO:
    """
    Generate an A7 landscape Word document for a single guest.

    Layout (vertically and horizontally centred):
        "Confirmacion de Asistencia"
        [QR code image]
        Name Surname

    A7 ISO 216: 74 mm x 105 mm (portrait) -> 105 mm x 74 mm (landscape)

    Returns:
        BytesIO buffer positioned at offset 0.
    """
    from docx import Document
    from docx.shared import Mm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Generate QR PNG
    qr_buf = generate_qr_code(rsvp_url, format="png", size=10)

    doc = Document()

    # Remove the default blank paragraph docx always inserts
    for para in doc.paragraphs:
        para._element.getparent().remove(para._element)

    # ------------------------------------------------------------------
    # Page: A7 landscape (ISO 216)
    # python-docx uses actual page dimensions; width > height = landscape.
    # ------------------------------------------------------------------
    section = doc.sections[0]
    section.page_width = Mm(105)   # long edge  = landscape width
    section.page_height = Mm(74)   # short edge = landscape height
    section.orientation = WD_ORIENT.LANDSCAPE

    # Margins: 6 mm all sides
    section.top_margin = Mm(6)
    section.bottom_margin = Mm(6)
    section.left_margin = Mm(6)
    section.right_margin = Mm(6)

    # Vertical centering via sectPr XML element
    sectPr = section._sectPr
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    sectPr.append(vAlign)

    # ------------------------------------------------------------------
    # Content
    # ------------------------------------------------------------------

    # 1. Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(4)
    title_run = title_para.add_run("Confirmaci\u00f3n de Asistencia")
    title_run.font.size = Pt(9)
    title_run.font.bold = True

    # 2. QR code (34 mm wide, centred)
    qr_para = doc.add_paragraph()
    qr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr_para.paragraph_format.space_before = Pt(2)
    qr_para.paragraph_format.space_after = Pt(4)
    qr_run = qr_para.add_run()
    qr_run.add_picture(qr_buf, width=Mm(34))

    # 3. Full name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(0)
    full_name = f"{name} {surname}" if surname else name
    name_run = name_para.add_run(full_name)
    name_run.font.size = Pt(9)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Routes: QR codes (with pending/confirmed filter)
# ---------------------------------------------------------------------------

@bp.route("/")
@admin_required
def qr_dashboard():
    """Display QR code dashboard - pending and confirmed guests only."""
    from app.models.guest import Guest

    all_guests = Guest.query.order_by(Guest.name).all()
    total_count = len(all_guests)
    guests = _filter_guests(all_guests)

    guest_data = []
    for guest in guests:
        rsvp_url = url_for("main.index", token=guest.token, _external=True)
        guest_data.append({
            "id":       guest.id,
            "name":     guest.name,
            "surname":  guest.surname,
            "phone":    guest.phone,
            "token":    guest.token,
            "rsvp_url": rsvp_url,
        })

    return render_template(
        "admin/qr_dashboard.html",
        guests=guest_data,
        total_count=total_count,
    )


@bp.route("/download/<int:guest_id>")
@bp.route("/download/<int:guest_id>/<format>")
@admin_required
def download_qr(guest_id, format="png"):
    """Download QR code image for a single guest (PNG or SVG)."""
    from app.models.guest import Guest

    guest = Guest.query.get_or_404(guest_id)
    rsvp_url = url_for("main.index", token=guest.token, _external=True)
    buf = generate_qr_code(rsvp_url, format=format)
    safe = _safe_filename(guest.name, guest.surname)

    if format == "svg":
        return send_file(buf, mimetype="image/svg+xml",
                         as_attachment=True, download_name=f"qr_{safe}.svg")
    return send_file(buf, mimetype="image/png",
                     as_attachment=True, download_name=f"qr_{safe}.png")


@bp.route("/preview/<int:guest_id>")
@admin_required
def preview_qr(guest_id):
    """Return QR code PNG inline (used for <img> tags in templates)."""
    from app.models.guest import Guest

    guest = Guest.query.get_or_404(guest_id)
    rsvp_url = url_for("main.index", token=guest.token, _external=True)
    buf = generate_qr_code(rsvp_url, format="png", size=8)
    return send_file(buf, mimetype="image/png", as_attachment=False)


@bp.route("/download-all")
@bp.route("/download-all/<format>")
@admin_required
def download_all_qr(format="png"):
    """Download a ZIP of QR images for pending and confirmed guests."""
    from app.models.guest import Guest

    guests = _filter_guests(Guest.query.order_by(Guest.name).all())
    zip_buf = BytesIO()

    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for guest in guests:
            rsvp_url = url_for("main.index", token=guest.token, _external=True)
            qr_buf = generate_qr_code(rsvp_url, format=format)
            safe = _safe_filename(guest.name, guest.surname)
            ext = "svg" if format == "svg" else "png"
            zf.writestr(f"{safe}.{ext}", qr_buf.getvalue())

    zip_buf.seek(0)
    return send_file(
        zip_buf,
        mimetype="application/zip",
        as_attachment=True,
        download_name=f"wedding_qr_codes_{format}.zip",
    )


@bp.route("/printable")
@admin_required
def printable_sheet():
    """Printable sheet of QR codes - pending and confirmed guests only."""
    from app.models.guest import Guest

    guests = _filter_guests(Guest.query.order_by(Guest.name).all())
    guest_data = []
    for guest in guests:
        rsvp_url = url_for("main.index", token=guest.token, _external=True)
        guest_data.append({
            "id":       guest.id,
            "name":     guest.name,
            "surname":  guest.surname,
            "rsvp_url": rsvp_url,
        })

    return render_template("admin/qr_printable.html", guests=guest_data)


# ---------------------------------------------------------------------------
# Routes: Individual DOCX download (synchronous — fast, ~0.3-0.5 s each)
# ---------------------------------------------------------------------------

@bp.route("/download-docx/<int:guest_id>")
@admin_required
def download_docx_single(guest_id):
    """
    Download an A7 landscape DOCX invitation card for one guest.
    Synchronous: generating a single document is fast and does not
    meaningfully block the server.
    """
    from app.models.guest import Guest

    guest = Guest.query.get_or_404(guest_id)
    rsvp_url = url_for("main.index", token=guest.token, _external=True)

    buf = _generate_guest_docx(guest.name, guest.surname, rsvp_url)
    safe = _safe_filename(guest.name, guest.surname)

    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"invitation_{safe}.docx",
    )


# ---------------------------------------------------------------------------
# Routes: Async bulk DOCX download
#
# Flow:
#   1. POST /admin/qr/download-all-docx/start
#          -> 200 {"job_id": "...", "total": N}
#   2. GET  /admin/qr/download-all-docx/status/<job_id>
#          -> 200 {"status": "running"|"done"|"error", "progress": 0-100, "total": N}
#   3. GET  /admin/qr/download-all-docx/download/<job_id>
#          -> 200 ZIP file  (only when status == "done")
#          -> 202 if not ready yet
# ---------------------------------------------------------------------------

@bp.route("/download-all-docx/start", methods=["POST"])
@admin_required
def start_download_all_docx():
    """
    Kick off background generation of a DOCX ZIP for all pending/confirmed guests.

    All database access happens here (within the Flask request context).
    The background thread only does CPU/IO work with plain Python dicts —
    no SQLAlchemy session, no Flask proxy objects.
    """
    from app.models.guest import Guest

    _cleanup_old_jobs()

    # Collect guest data while we have the app/request context
    all_guests = Guest.query.order_by(Guest.name).all()
    guests_data = []
    for g in _filter_guests(all_guests):
        rsvp_url = url_for("main.index", token=g.token, _external=True)
        guests_data.append({
            "name":      g.name,
            "surname":   g.surname,
            "rsvp_url":  rsvp_url,
            "safe_name": _safe_filename(g.name, g.surname),
        })

    job_id = str(uuid.uuid4())
    with _jobs_lock:
        _docx_jobs[job_id] = {
            "status":     "running",
            "progress":   0,
            "total":      len(guests_data),
            "data":       None,
            "error":      None,
            "created_at": datetime.utcnow(),
        }

    def _generate_zip() -> None:
        """Background thread: build ZIP, update job state. No Flask context needed."""
        try:
            zip_buf = BytesIO()
            total = len(guests_data)

            with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                for i, info in enumerate(guests_data):
                    docx_buf = _generate_guest_docx(
                        info["name"], info["surname"], info["rsvp_url"]
                    )
                    zf.writestr(f"{info['safe_name']}.docx", docx_buf.getvalue())

                    progress = int((i + 1) / total * 100)
                    with _jobs_lock:
                        _docx_jobs[job_id]["progress"] = progress

            zip_buf.seek(0)
            with _jobs_lock:
                _docx_jobs[job_id]["status"] = "done"
                _docx_jobs[job_id]["data"]   = zip_buf

            logger.info(f"DOCX ZIP job {job_id} completed ({total} files).")

        except Exception as exc:
            logger.error(f"DOCX ZIP job {job_id} failed: {exc}", exc_info=True)
            with _jobs_lock:
                _docx_jobs[job_id]["status"] = "error"
                _docx_jobs[job_id]["error"]  = str(exc)

    thread = threading.Thread(
        target=_generate_zip,
        daemon=True,
        name=f"docx-zip-{job_id[:8]}",
    )
    thread.start()

    return jsonify({"job_id": job_id, "total": len(guests_data)})


@bp.route("/download-all-docx/status/<job_id>")
@admin_required
def download_all_docx_status(job_id):
    """Poll the progress of an async DOCX ZIP job."""
    with _jobs_lock:
        job = _docx_jobs.get(job_id)

    if not job:
        return jsonify({"error": "Job not found or expired"}), 404

    return jsonify({
        "status":   job["status"],
        "progress": job["progress"],
        "total":    job["total"],
        "error":    job.get("error"),
    })


@bp.route("/download-all-docx/download/<job_id>")
@admin_required
def download_all_docx_file(job_id):
    """Download the completed DOCX ZIP. Removes the job from memory after serving."""
    with _jobs_lock:
        job = _docx_jobs.get(job_id)

    if not job:
        return jsonify({"error": "Job not found or expired"}), 404
    if job["status"] != "done":
        return jsonify({"error": "Job not ready yet", "status": job["status"]}), 202

    data = job["data"]
    data.seek(0)

    with _jobs_lock:
        _docx_jobs.pop(job_id, None)

    return send_file(
        data,
        mimetype="application/zip",
        as_attachment=True,
        download_name="wedding_invitations_docx.zip",
    )