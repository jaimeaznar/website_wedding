# tests/test_pdf_service.py
import pytest
import io
from docx import Document
from app import db
from app.models.guest import Guest
from app.models.rsvp import RSVP, AdditionalGuest
from app.models.allergen import Allergen, GuestAllergen
from app.services.pdf_service import PDFService
from PyPDF2 import PdfReader


class TestPDFService:
    """Test cases for PDF generation service."""
    
    @pytest.fixture
    def sample_data(self, app):
        """Create sample data for PDF testing."""
        with app.app_context():
            # Create allergens (check if they exist first)
            allergens = []
            for name in ['Gluten', 'Dairy', 'Nuts', 'Shellfish']:
                allergen = Allergen.query.filter_by(name=name).first()
                if not allergen:
                    allergen = Allergen(name=name)
                    db.session.add(allergen)
                allergens.append(allergen)
            db.session.flush()
            
            # Create guests with RSVPs and allergens
            guests_data = [
                {
                    'name': 'John Doe',
                    'phone': '555-0001',
                    'hotel': 'Hotel Parador Trujillo',
                    'transport_church': True,
                    'transport_reception': True,
                    'transport_hotel': True,
                    'allergens': [allergens[0], allergens[1]]  # Gluten, Dairy
                },
                {
                    'name': 'Jane Smith',
                    'phone': '555-0002',
                    'hotel': 'Hotel Cáceres',
                    'transport_church': True,
                    'transport_reception': False,
                    'transport_hotel': True,
                    'allergens': [allergens[2]]  # Nuts
                },
                {
                    'name': 'Bob Johnson',
                    'phone': '555-0003',
                    'hotel': 'Hotel Parador Trujillo',
                    'transport_church': False,
                    'transport_reception': True,
                    'transport_hotel': False,
                    'allergens': []
                }
            ]
            
            created_guests = []
            for data in guests_data:
                guest = Guest(
                    name=data['name'],
                    phone=data['phone'],
                    token=f"token-{data['phone']}"
                )
                db.session.add(guest)
                db.session.flush()
                
                rsvp = RSVP(
                    guest_id=guest.id,
                    is_attending=True,
                    hotel_name=data['hotel'],
                    transport_to_reception=data['transport_reception'],
                    transport_to_hotel=data['transport_hotel']
                )
                db.session.add(rsvp)
                db.session.flush()
                
                # Add allergens
                for allergen in data['allergens']:
                    guest_allergen = GuestAllergen(
                        rsvp_id=rsvp.id,
                        guest_name=guest.name,
                        allergen_id=allergen.id
                    )
                    db.session.add(guest_allergen)
                
                # Add custom allergen for first guest
                if guest.name == 'John Doe':
                    custom_allergen = GuestAllergen(
                        rsvp_id=rsvp.id,
                        guest_name=guest.name,
                        custom_allergen='Vegetarian'
                    )
                    db.session.add(custom_allergen)
                
                created_guests.append(guest)
            
            db.session.commit()
            
            yield created_guests
            
            # Cleanup
            for guest in created_guests:
                if Guest.query.get(guest.id):
                    db.session.delete(guest)
            db.session.commit()
    
    def test_generate_dietary_pdf_returns_bytes(self, app, sample_data):
        """Test that dietary PDF generation returns bytes."""
        with app.app_context():
            pdf_data = PDFService.generate_dietary_pdf()
            
            assert isinstance(pdf_data, bytes)
            assert len(pdf_data) > 0
    
    def test_dietary_pdf_is_valid_pdf(self, app, sample_data):
        """Test that generated dietary PDF is a valid PDF file."""
        with app.app_context():
            pdf_data = PDFService.generate_dietary_pdf()
            
            # Try to read PDF
            pdf_buffer = io.BytesIO(pdf_data)
            reader = PdfReader(pdf_buffer)
            
            assert len(reader.pages) > 0
    
    def test_dietary_pdf_contains_guest_names(self, app, sample_data):
        """Test that dietary PDF contains guest names."""
        with app.app_context():
            pdf_data = PDFService.generate_dietary_pdf()
            
            pdf_buffer = io.BytesIO(pdf_data)
            reader = PdfReader(pdf_buffer)
            
            # Extract text from all pages
            text = ''
            for page in reader.pages:
                text += page.extract_text()
            
            # Check for guest names
            assert 'John Doe' in text
            assert 'Jane Smith' in text
    
    def test_dietary_pdf_contains_allergens(self, app, sample_data):
        """Test that dietary PDF contains allergen information."""
        with app.app_context():
            pdf_data = PDFService.generate_dietary_pdf()
            
            pdf_buffer = io.BytesIO(pdf_data)
            reader = PdfReader(pdf_buffer)
            
            text = ''
            for page in reader.pages:
                text += page.extract_text()
            
            # Check for allergens
            assert 'Gluten' in text
            assert 'Dairy' in text
            assert 'Nuts' in text
    
    def test_dietary_pdf_contains_custom_allergen(self, app, sample_data):
        """Test that dietary PDF includes custom allergens."""
        with app.app_context():
            pdf_data = PDFService.generate_dietary_pdf()
            
            pdf_buffer = io.BytesIO(pdf_data)
            reader = PdfReader(pdf_buffer)
            
            text = ''
            for page in reader.pages:
                text += page.extract_text()
            
            assert 'Vegetarian' in text
    
    def test_dietary_pdf_has_header(self, app, sample_data):
        """Test that dietary PDF contains wedding header."""
        with app.app_context():
            pdf_data = PDFService.generate_dietary_pdf()
            
            pdf_buffer = io.BytesIO(pdf_data)
            reader = PdfReader(pdf_buffer)
            
            text = reader.pages[0].extract_text()
            
            assert "Irene & Jaime" in text or "Wedding" in text
    
    def test_generate_transport_pdf_returns_bytes(self, app, sample_data):
        """Test that transport PDF generation returns bytes."""
        with app.app_context():
            pdf_data = PDFService.generate_transport_pdf()
            
            assert isinstance(pdf_data, bytes)
            assert len(pdf_data) > 0
    
    def test_transport_pdf_is_valid_pdf(self, app, sample_data):
        """Test that generated transport PDF is a valid PDF file."""
        with app.app_context():
            pdf_data = PDFService.generate_transport_pdf()
            
            pdf_buffer = io.BytesIO(pdf_data)
            reader = PdfReader(pdf_buffer)
            
            assert len(reader.pages) > 0
    
    def test_transport_pdf_contains_guest_names(self, app, sample_data):
        """Test that transport PDF contains guest names."""
        with app.app_context():
            pdf_data = PDFService.generate_transport_pdf()
            
            pdf_buffer = io.BytesIO(pdf_data)
            reader = PdfReader(pdf_buffer)
            
            text = ''
            for page in reader.pages:
                text += page.extract_text()
            
            assert 'John Doe' in text
            assert 'Jane Smith' in text
    
    def test_transport_pdf_contains_hotel_names(self, app, sample_data):
        """Test that transport PDF contains hotel information."""
        with app.app_context():
            pdf_data = PDFService.generate_transport_pdf()
            
            pdf_buffer = io.BytesIO(pdf_data)
            reader = PdfReader(pdf_buffer)
            
            text = ''
            for page in reader.pages:
                text += page.extract_text()
            
            assert 'Hotel Parador Trujillo' in text or 'Parador Trujillo' in text
            assert 'Cáceres' in text
    
    def test_transport_pdf_contains_route_information(self, app, sample_data):
        """Test that transport PDF contains route descriptions."""
        with app.app_context():
            pdf_data = PDFService.generate_transport_pdf()
            
            pdf_buffer = io.BytesIO(pdf_data)
            reader = PdfReader(pdf_buffer)
            
            text = ''
            for page in reader.pages:
                text += page.extract_text()
            
            # Check for route descriptions
            assert 'Church' in text
            assert 'Reception' in text
            assert 'Hotel' in text
    
    def test_transport_pdf_contains_contact_info(self, app, sample_data):
        """Test that transport PDF includes phone numbers."""
        with app.app_context():
            pdf_data = PDFService.generate_transport_pdf()
            
            pdf_buffer = io.BytesIO(pdf_data)
            reader = PdfReader(pdf_buffer)
            
            text = ''
            for page in reader.pages:
                text += page.extract_text()
            
            # Check for phone numbers
            assert '555-0001' in text or '5550001' in text
    
    def test_transport_pdf_groups_by_hotel(self, app, sample_data):
        """Test that transport PDF groups guests by hotel."""
        with app.app_context():
            pdf_data = PDFService.generate_transport_pdf()
            
            pdf_buffer = io.BytesIO(pdf_data)
            reader = PdfReader(pdf_buffer)
            
            text = ''
            for page in reader.pages:
                text += page.extract_text()
            
            # Both hotels should appear
            assert 'Parador' in text or 'Trujillo' in text
            assert 'Cáceres' in text
    
    def test_pdf_generation_with_no_data(self, app):
        """Test PDF generation when no RSVP data exists."""
        with app.app_context():
            # Should not crash, should generate empty/minimal PDF
            pdf_data = PDFService.generate_dietary_pdf()
            assert isinstance(pdf_data, bytes)
            assert len(pdf_data) > 0
            
            pdf_data = PDFService.generate_transport_pdf()
            assert isinstance(pdf_data, bytes)
            assert len(pdf_data) > 0
    
    def test_dietary_pdf_with_many_allergens(self, app):
        """Test dietary PDF generation with multiple allergens per guest."""
        with app.app_context():
            # Create guest with many allergens
            allergens = []
            for name in ['Gluten', 'Dairy', 'Nuts', 'Shellfish', 'Eggs', 'Soy']:
                allergen = Allergen.query.filter_by(name=name).first()
                if not allergen:
                    allergen = Allergen(name=name)
                    db.session.add(allergen)
                allergens.append(allergen)
            db.session.flush()
            
            guest = Guest(
                name='Allergy Test Guest',
                phone='555-9999',
                token='token-allergy'
            )
            db.session.add(guest)
            db.session.flush()
            
            rsvp = RSVP(
                guest_id=guest.id,
                is_attending=True,
                hotel_name='Test Hotel'
            )
            db.session.add(rsvp)
            db.session.flush()
            
            # Add multiple allergens
            for allergen in allergens:
                guest_allergen = GuestAllergen(
                    rsvp_id=rsvp.id,
                    guest_name=guest.name,
                    allergen_id=allergen.id
                )
                db.session.add(guest_allergen)
            
            db.session.commit()
            
            # Generate PDF
            pdf_data = PDFService.generate_dietary_pdf()
            
            assert isinstance(pdf_data, bytes)
            assert len(pdf_data) > 0
            
            # Verify PDF is valid
            pdf_buffer = io.BytesIO(pdf_data)
            reader = PdfReader(pdf_buffer)
            assert len(reader.pages) > 0
            
            # Cleanup
            db.session.delete(guest)
            db.session.commit()
    
    def test_transport_pdf_with_additional_guests(self, app):
        """Test transport PDF includes additional guests in count."""
        with app.app_context():
            guest = Guest(
                name='Family Test',
                phone='555-8888',

                token='token-family',

            )
            db.session.add(guest)
            db.session.flush()
            
            rsvp = RSVP(
                guest_id=guest.id,
                is_attending=True,
                hotel_name='Family Hotel',
                transport_to_reception=True,
                transport_to_hotel=True
            )
            db.session.add(rsvp)
            db.session.flush()
            
            # Add additional guests
            additional1 = AdditionalGuest(
                rsvp_id=rsvp.id,
                name='Spouse',
                is_child=False
            )
            additional2 = AdditionalGuest(
                rsvp_id=rsvp.id,
                name='Child',
                is_child=True
            )
            db.session.add(additional1)
            db.session.add(additional2)
            db.session.commit()
            
            # Generate PDF
            pdf_data = PDFService.generate_transport_pdf()
            
            pdf_buffer = io.BytesIO(pdf_data)
            reader = PdfReader(pdf_buffer)
            
            text = ''
            for page in reader.pages:
                text += page.extract_text()
            
            # Should show count of 3 (main + 2 additional)
            assert 'Family Test' in text
            assert '3' in text  # Total passenger count
            
            # Cleanup
            db.session.delete(guest)
            db.session.commit()
    
    def test_pdf_service_page_count(self, app, sample_data):
        """Test that PDFs have multiple pages when needed."""
        with app.app_context():
            dietary_pdf = PDFService.generate_dietary_pdf()
            transport_pdf = PDFService.generate_transport_pdf()
            
            dietary_reader = PdfReader(io.BytesIO(dietary_pdf))
            transport_reader = PdfReader(io.BytesIO(transport_pdf))
            
            # Should have at least 1 page each
            assert len(dietary_reader.pages) >= 1
            assert len(transport_reader.pages) >= 1
    
    def test_pdf_headers_and_footers(self, app, sample_data):
        """Test that PDFs include headers and footers."""
        with app.app_context():
            pdf_data = PDFService.generate_dietary_pdf()
            
            pdf_buffer = io.BytesIO(pdf_data)
            reader = PdfReader(pdf_buffer)
            
            # Check first page
            text = reader.pages[0].extract_text()
            
            # Should have wedding names in header
            assert 'Irene' in text or 'Jaime' in text or 'Wedding' in text
            
            # Should have generation date
            assert 'Generated' in text or '202' in text  # Year format

class TestDietaryDocx:
    """Tests for generate_dietary_docx — the catering Word document with one
    table row per attendee (main guests + plus-ones grouped together)."""

    @pytest.fixture
    def docx_sample(self, app):
        """Setup covering every branch:
          - Ana: main guest with 2 plus-ones (Luis, Sofía) -> "(acomp. N)" labels
          - Juan: main guest with 1 plus-one (María) -> "(acompañante)" label
          - Pedro: solo attendee, no restrictions
          - Carmen: cancelled (must be excluded)
          - Mario: declined (must be excluded)
        """
        with app.app_context():
            allergens = {}
            for name in ['Gluten', 'Frutos secos']:
                a = Allergen.query.filter_by(name=name).first()
                if not a:
                    a = Allergen(name=name)
                    db.session.add(a)
                allergens[name] = a
            db.session.flush()

            created = []

            # 1) Ana + 2 plus-ones; Sofía also has a custom restriction
            ana = Guest(name='Ana', surname='Martínez Ruiz',
                        phone='555-d1', token='tok-docx-1')
            db.session.add(ana)
            db.session.flush()
            ana_rsvp = RSVP(guest_id=ana.id, is_attending=True)
            db.session.add(ana_rsvp)
            db.session.flush()
            db.session.add(GuestAllergen(rsvp_id=ana_rsvp.id, guest_name='Ana',
                                         allergen_id=allergens['Frutos secos'].id))
            db.session.add(GuestAllergen(rsvp_id=ana_rsvp.id, guest_name='Ana',
                                         custom_allergen='Vegana'))
            luis = AdditionalGuest(rsvp_id=ana_rsvp.id, name='Luis')
            sofia = AdditionalGuest(rsvp_id=ana_rsvp.id, name='Sofía')
            db.session.add_all([luis, sofia])
            db.session.flush()
            db.session.add(GuestAllergen(rsvp_id=ana_rsvp.id, guest_name='Sofía',
                                         custom_allergen='Lactosa'))
            created.append(ana)

            # 2) Juan + 1 plus-one
            juan = Guest(name='Juan', surname='García López',
                         phone='555-d2', token='tok-docx-2')
            db.session.add(juan)
            db.session.flush()
            juan_rsvp = RSVP(guest_id=juan.id, is_attending=True)
            db.session.add(juan_rsvp)
            db.session.flush()
            db.session.add(GuestAllergen(rsvp_id=juan_rsvp.id, guest_name='Juan',
                                         allergen_id=allergens['Gluten'].id))
            db.session.add(AdditionalGuest(rsvp_id=juan_rsvp.id, name='María'))
            created.append(juan)

            # 3) Pedro - solo, no restrictions
            pedro = Guest(name='Pedro', surname='Sánchez Ruiz',
                          phone='555-d3', token='tok-docx-3')
            db.session.add(pedro)
            db.session.flush()
            db.session.add(RSVP(guest_id=pedro.id, is_attending=True))
            created.append(pedro)

            # 4) Carmen - cancelled, must NOT appear
            carmen = Guest(name='Carmen', surname='Cancelada',
                           phone='555-d4', token='tok-docx-4')
            db.session.add(carmen)
            db.session.flush()
            db.session.add(RSVP(guest_id=carmen.id, is_attending=False,
                                is_cancelled=True))
            created.append(carmen)

            # 5) Mario - declined, must NOT appear
            mario = Guest(name='Mario', surname='Declinado',
                          phone='555-d5', token='tok-docx-5')
            db.session.add(mario)
            db.session.flush()
            db.session.add(RSVP(guest_id=mario.id, is_attending=False))
            created.append(mario)

            db.session.commit()

            yield created

            for g in created:
                if Guest.query.get(g.id):
                    db.session.delete(g)
            db.session.commit()

    # ----- helpers ---------------------------------------------------------

    @staticmethod
    def _open(docx_bytes):
        return Document(io.BytesIO(docx_bytes))

    @staticmethod
    def _data_rows(table):
        """Return list of (nombre, apellidos, mesa, restricciones) for non-header rows."""
        return [
            (r.cells[0].text, r.cells[1].text, r.cells[2].text, r.cells[3].text)
            for r in table.rows[1:]
        ]

    # ----- tests -----------------------------------------------------------

    def test_returns_valid_docx_bytes(self, app, docx_sample):
        with app.app_context():
            data = PDFService.generate_dietary_docx()
            assert isinstance(data, bytes) and len(data) > 0
            doc = self._open(data)
            assert len(doc.tables) >= 1

    def test_table_has_four_columns_and_correct_headers(self, app, docx_sample):
        with app.app_context():
            doc = self._open(PDFService.generate_dietary_docx())
            table = doc.tables[0]
            assert len(table.columns) == 4
            header = [c.text for c in table.rows[0].cells]
            assert header == ['Nombre', 'Apellidos', 'Mesa', 'Restricciones']

    def test_main_guests_appear_with_their_surnames(self, app, docx_sample):
        with app.app_context():
            doc = self._open(PDFService.generate_dietary_docx())
            rows = self._data_rows(doc.tables[0])
            name_surname = {(n, s) for n, s, _, _ in rows}
            assert ('Ana', 'Martínez Ruiz') in name_surname
            assert ('Juan', 'García López') in name_surname
            assert ('Pedro', 'Sánchez Ruiz') in name_surname

    def test_excludes_cancelled_and_declined_guests(self, app, docx_sample):
        with app.app_context():
            doc = self._open(PDFService.generate_dietary_docx())
            all_text = ' '.join(c.text for r in doc.tables[0].rows for c in r.cells)
            assert 'Carmen' not in all_text
            assert 'Cancelada' not in all_text
            assert 'Mario' not in all_text
            assert 'Declinado' not in all_text

    def test_plus_ones_appear_immediately_after_main_guest_in_order(self, app, docx_sample):
        """Family unit must stay together so the catering team reads one party at a time."""
        with app.app_context():
            doc = self._open(PDFService.generate_dietary_docx())
            names = [r.cells[0].text for r in doc.tables[0].rows[1:]]
            ana_idx = names.index('Ana')
            assert names[ana_idx + 1].startswith('Luis')
            assert names[ana_idx + 2].startswith('Sofía')
            juan_idx = names.index('Juan')
            assert names[juan_idx + 1].startswith('María')

    def test_single_plus_one_uses_acompanante_label(self, app, docx_sample):
        with app.app_context():
            doc = self._open(PDFService.generate_dietary_docx())
            names = [r.cells[0].text for r in doc.tables[0].rows[1:]]
            maria_rows = [n for n in names if n.startswith('María')]
            assert len(maria_rows) == 1
            assert '(acompañante)' in maria_rows[0]

    def test_multiple_plus_ones_use_numbered_labels(self, app, docx_sample):
        with app.app_context():
            doc = self._open(PDFService.generate_dietary_docx())
            names = [r.cells[0].text for r in doc.tables[0].rows[1:]]
            assert any(n.startswith('Luis') and 'acomp. 1' in n for n in names)
            assert any(n.startswith('Sofía') and 'acomp. 2' in n for n in names)

    def test_plus_ones_inherit_main_guests_surname(self, app, docx_sample):
        """Surname column shows the family surname so each party visually clusters."""
        with app.app_context():
            doc = self._open(PDFService.generate_dietary_docx())
            rows = self._data_rows(doc.tables[0])
            surname_for = {n.split(' (')[0]: s for n, s, _, _ in rows}
            assert surname_for['Luis'] == 'Martínez Ruiz'
            assert surname_for['Sofía'] == 'Martínez Ruiz'
            assert surname_for['María'] == 'García López'

    def test_mesa_column_is_blank_for_every_data_row(self, app, docx_sample):
        with app.app_context():
            doc = self._open(PDFService.generate_dietary_docx())
            for row in doc.tables[0].rows[1:]:
                assert row.cells[2].text == ''

    def test_restrictions_appear_against_the_correct_person(self, app, docx_sample):
        with app.app_context():
            doc = self._open(PDFService.generate_dietary_docx())
            rows = self._data_rows(doc.tables[0])
            restr_for = {n.split(' (')[0]: r for n, _, _, r in rows}

            assert 'Frutos secos' in restr_for['Ana']
            assert 'Vegana' in restr_for['Ana']
            assert 'Gluten' in restr_for['Juan']
            assert 'Lactosa' in restr_for['Sofía']
            # Plus-ones with no restrictions and main guests with no restrictions
            # should not silently inherit anyone else's allergens
            assert restr_for['Pedro'] == ''
            assert restr_for['Luis'] == ''
            assert restr_for['María'] == ''

    def test_total_attendee_rows(self, app, docx_sample):
        """Header + Ana + Luis + Sofía + Juan + María + Pedro = 7 rows."""
        with app.app_context():
            doc = self._open(PDFService.generate_dietary_docx())
            assert len(doc.tables[0].rows) == 7

    def test_does_not_crash_when_no_attendees(self, app):
        """Empty database should yield a valid (header-only) document, not an exception."""
        with app.app_context():
            data = PDFService.generate_dietary_docx()
            assert isinstance(data, bytes)
            doc = self._open(data)
            assert len(doc.tables) >= 1
            # At minimum the header row exists
            assert len(doc.tables[0].rows) >= 1